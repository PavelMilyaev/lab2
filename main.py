import os
import sys
from clothing_package import Jacket, Trousers, ThreePieceSuit, ClothingCalculator
from database import init_db, save_calculation, get_history

try:
    from docx import Document
    from datetime import datetime

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docx не установлен. Сохранение в .docx недоступно.")


class ClothingApp:
    """Главное приложение"""

    def __init__(self):
        self.calculator = ClothingCalculator()
        init_db()

    def clear_screen(self):
        """Очистка экрана (безопасная версия)"""
        try:
            # Проверяем, есть ли поддержка терминала
            if sys.stdout.isatty():
                os.system('cls' if os.name == 'nt' else 'clear')
            else:
                # Если нет терминала, просто печатаем разделитель
                print("\n")
        except:
            # В случае любой ошибки, печатаем разделитель
            print("\n")

    def show_menu(self):
        """Показ главного меню"""
        self.clear_screen()
        print("КАЛЬКУЛЯТОР РАСХОДА ТКАНИ И СТОИМОСТИ ПОШИВА")
        print("1. Добавить пиджак")
        print("2. Добавить брюки")
        print("3. Создать костюм-тройку")
        print("4. Показать все изделия")
        print("5. Удалить изделие")
        print("6. Рассчитать итоги")
        if DOCX_AVAILABLE:
            print("7. Сохранить отчёт в .docx")
        print("8. Показать историю расчётов")
        print("9. Выход")

    def get_input(self, prompt, input_type=str, validation=None):
        """Безопасный ввод данных с обработкой ошибок"""
        while True:
            try:
                value = input(prompt)
                if not value.strip():  # Пустой ввод
                    print("Ошибка: введите значение")
                    continue

                converted_value = input_type(value)

                if validation and not validation(converted_value):
                    print(f"Ошибка: значение должно удовлетворять условию {validation.__doc__}")
                    continue

                return converted_value
            except ValueError:
                print(f"Ошибка: введите значение типа {input_type.__name__}")
            except KeyboardInterrupt:
                print("\nВыход из программы...")
                sys.exit(0)
            except EOFError:
                print("\nОбнаружен конец файла. Выход...")
                sys.exit(0)

    def validate_size(self, size):
        """Проверка размера (44-64)"""
        return 44 <= size <= 64

    def validate_positive(self, value):
        """Проверка положительного числа"""
        return value >= 0

    def add_jacket(self):
        """Добавление пиджака"""
        self.clear_screen()
        print("ДОБАВЛЕНИЕ ПИДЖАКА")

        try:
            name = self.get_input("Название модели: ", str)
            size = self.get_input("Размер (44-64): ", int, self.validate_size)
            fabric_price = self.get_input("Цена ткани за метр (руб): ", float, self.validate_positive)
            accessories_price = self.get_input("Стоимость фурнитуры (руб): ", float, self.validate_positive)

            has_lining = self.get_input("Есть подкладка? (д/н): ", str).lower() == 'д'
            pockets = self.get_input("Количество карманов: ", int, self.validate_positive)

            jacket = Jacket(name, size, fabric_price, accessories_price, has_lining, pockets)
            self.calculator.add_item(jacket)

            print(f"\nПиджак '{name}' добавлен!")
            print(f"Расход ткани: {jacket.calculate_fabric_consumption()} м")
            print(f"Стоимость пошива: {jacket.calculate_sewing_cost()} руб")
            print(f"Общая стоимость: {jacket.calculate_total_cost()} руб")

        except Exception as e:
            print(f"\nОшибка при добавлении пиджака: {e}")

        self.wait_for_enter()

    def add_trousers(self):
        """Добавление брюк"""
        self.clear_screen()
        print("ДОБАВЛЕНИЕ БРЮК")

        try:
            name = self.get_input("Название модели: ", str)
            size = self.get_input("Размер (44-64): ", int, self.validate_size)
            fabric_price = self.get_input("Цена ткани за метр (руб): ", float, self.validate_positive)
            accessories_price = self.get_input("Стоимость фурнитуры (руб): ", float, self.validate_positive)

            has_belt = self.get_input("Есть пояс? (д/н): ", str).lower() == 'д'
            is_classic = self.get_input("Классические? (д/н): ", str).lower() == 'д'

            trousers = Trousers(name, size, fabric_price, accessories_price, has_belt, is_classic)
            self.calculator.add_item(trousers)

            print(f"\nБрюки '{name}' добавлены!")
            print(f"Расход ткани: {trousers.calculate_fabric_consumption()} м")
            print(f"Стоимость пошива: {trousers.calculate_sewing_cost()} руб")
            print(f"Общая стоимость: {trousers.calculate_total_cost()} руб")

        except Exception as e:
            print(f"\nОшибка при добавлении брюк: {e}")

        self.wait_for_enter()

    def create_suit(self):
        """Создание костюма из существующих элементов"""
        self.clear_screen()
        print("СОЗДАНИЕ КОСТЮМА-ТРОЙКИ")

        if len(self.calculator) < 2:
            print("Недостаточно изделий для создания костюма!")
            self.wait_for_enter()
            return

        print("\nДоступные изделия:")
        items = self.calculator.get_items()

        # Показываем только пиджаки и брюки
        jackets = []
        trousers_list = []

        for i, item in enumerate(items):
            if isinstance(item, Jacket):
                jackets.append((i, item))
                print(f"{i + 1}. {item} (Пиджак)")
            elif isinstance(item, Trousers):
                trousers_list.append((i, item))
                print(f"{i + 1}. {item} (Брюки)")

        if not jackets or not trousers_list:
            print("\nНе хватает пиджака или брюк для создания костюма!")
            self.wait_for_enter()
            return

        try:
            print("\nВыберите пиджак:")
            jacket_idx = self.get_input("Номер пиджака: ", int) - 1

            print("\nВыберите брюки:")
            trousers_idx = self.get_input("Номер брюк: ", int) - 1

            # Проверяем, что выбранные индексы соответствуют пиджаку и брюкам
            selected_jacket = None
            selected_trousers = None

            for idx, item in jackets:
                if idx == jacket_idx:
                    selected_jacket = item
                    break

            for idx, item in trousers_list:
                if idx == trousers_idx:
                    selected_trousers = item
                    break

            if not selected_jacket or not selected_trousers:
                print("Ошибка: выберите корректные номера пиджака и брюк")
                self.wait_for_enter()
                return

            name = self.get_input("Название костюма: ", str)
            suit = ThreePieceSuit(name, selected_jacket, selected_trousers, None)
            self.calculator.add_item(suit)

            print(f"\nКостюм '{name}' создан!")
            print(f"Общий расход ткани: {suit.calculate_fabric_consumption()} м")
            print(f"Стоимость пошива (со скидкой): {suit.calculate_sewing_cost()} руб")
            print(f"Общая стоимость материалов: {suit.calculate_total_cost()} руб")

        except Exception as e:
            print(f"\nОшибка при создании костюма: {e}")

        self.wait_for_enter()

    def show_items(self):
        """Показ всех изделий"""
        self.clear_screen()
        print("ВСЕ ИЗДЕЛИЯ")

        items = self.calculator.get_items()
        if not items:
            print("Нет добавленных изделий")
        else:
            for i, item in enumerate(items):
                print(f"\n{i + 1}. {item}")
                print(f"   Тип: {item.__class__.__name__}")
                print(f"   Расход ткани: {item.calculate_fabric_consumption()} м")
                print(f"   Стоимость пошива: {item.calculate_sewing_cost()} руб")
                print(f"   Стоимость материалов: {item.calculate_total_cost()} руб")
                print(f"   Общая стоимость: {item.calculate_total_cost() + item.calculate_sewing_cost()} руб")

        self.wait_for_enter()

    def remove_item(self):
        """Удаление изделия"""
        self.clear_screen()
        print("УДАЛЕНИЕ ИЗДЕЛИЯ")

        items = self.calculator.get_items()
        if not items:
            print("Нет изделий для удаления")
            self.wait_for_enter()
            return

        for i, item in enumerate(items):
            print(f"{i + 1}. {item}")

        try:
            idx = self.get_input("Выберите номер для удаления (0 для отмены): ", int) - 1
            if idx >= 0:
                self.calculator.remove_item(idx)
                print("Изделие удалено!")
            else:
                print("Удаление отменено")
        except Exception as e:
            print(f"Ошибка при удалении: {e}")

        self.wait_for_enter()

    def calculate_totals(self):
        """Расчёт итогов и сохранение в БД"""
        self.clear_screen()
        print("ИТОГОВЫЙ РАСЧЁТ")

        if len(self.calculator) == 0:
            print("Нет изделий для расчёта")
            self.wait_for_enter()
            return

        total_fabric = self.calculator.calculate_total_fabric()
        total_sewing = self.calculator.calculate_total_sewing_cost()
        total_materials = self.calculator.calculate_total_material_cost()
        grand_total = total_sewing + total_materials

        print(f"Количество изделий: {len(self.calculator)}")
        print(f"Общий расход ткани: {total_fabric:.2f} м")
        print(f"Общая стоимость пошива: {total_sewing:.2f} руб")
        print(f"Общая стоимость материалов: {total_materials:.2f} руб")
        print(f"ИТОГО: {grand_total:.2f} руб")

        # Сохраняем в БД
        try:
            details = "\n".join([f"{i + 1}. {item}" for i, item in enumerate(self.calculator.get_items())])
            save_calculation(details, total_fabric, total_sewing, total_materials, grand_total)
            print("\nРасчёт сохранён в базу данных")
        except Exception as e:
            print(f"\nОшибка при сохранении в БД: {e}")

        self.wait_for_enter()

    def save_report(self):
        """Сохранение отчёта в .docx"""
        if not DOCX_AVAILABLE:
            print("Библиотека python-docx не установлена.")
            print("Установите: pip install python-docx")
            self.wait_for_enter()
            return

        self.clear_screen()
        print("СОХРАНЕНИЕ ОТЧЁТА")

        if len(self.calculator) == 0:
            print("Нет данных для сохранения")
            self.wait_for_enter()
            return

        try:
            doc = Document()
            doc.add_heading('Отчёт о расчёте одежды', 0)

            # Дата
            doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")

            doc.add_heading('Изделия:', level=1)
            for i, item in enumerate(self.calculator.get_items()):
                doc.add_heading(f"{i + 1}. {item}", level=2)
                doc.add_paragraph(f"Тип: {item.__class__.__name__}", style='List Bullet')
                doc.add_paragraph(f"Расход ткани: {item.calculate_fabric_consumption()} м", style='List Bullet')
                doc.add_paragraph(f"Стоимость пошива: {item.calculate_sewing_cost()} руб", style='List Bullet')
                doc.add_paragraph(f"Стоимость материалов: {item.calculate_total_cost()} руб", style='List Bullet')

            doc.add_heading('Итоги:', level=1)
            doc.add_paragraph(f"Общий расход ткани: {self.calculator.calculate_total_fabric():.2f} м")
            doc.add_paragraph(f"Общая стоимость пошива: {self.calculator.calculate_total_sewing_cost():.2f} руб")
            doc.add_paragraph(f"Общая стоимость материалов: {self.calculator.calculate_total_material_cost():.2f} руб")
            doc.add_paragraph(
                f"ИТОГО: {self.calculator.calculate_total_sewing_cost() + self.calculator.calculate_total_material_cost():.2f} руб")

            filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(filename)
            print(f"\nОтчёт сохранён как {filename}")
            print(f"   Полный путь: {os.path.abspath(filename)}")

        except Exception as e:
            print(f"\nОшибка при сохранении отчёта: {e}")

        self.wait_for_enter()

    def show_history(self):
        """Показ истории расчётов из БД"""
        self.clear_screen()
        print("ИСТОРИЯ РАСЧЁТОВ")

        try:
            history = get_history()
            if not history:
                print("История пуста")
            else:
                for record in history:
                    print(f"ID: {record[0]} | Дата: {record[1]}")
                    print("Изделия:")
                    for line in record[2].split('\n'):
                        print(f"  {line}")
                    print(f"\nРасход ткани: {record[3]:.2f} м")
                    print(f"Стоимость пошива: {record[4]:.2f} руб")
                    print(f"Стоимость материалов: {record[5]:.2f} руб")
                    print(f"ИТОГО: {record[6]:.2f} руб")
        except Exception as e:
            print(f"Ошибка при загрузке истории: {e}")

        self.wait_for_enter()

    def wait_for_enter(self):
        """Ожидание нажатия Enter (безопасная версия)"""
        try:
            input("\nНажмите Enter для продолжения...")
        except (KeyboardInterrupt, EOFError):
            print("\nВозврат в меню...")

    def run(self):
        """Запуск приложения"""

        while True:
            try:
                self.show_menu()
                choice = self.get_input("Выберите действие: ", str)

                if choice == '1':
                    self.add_jacket()
                elif choice == '2':
                    self.add_trousers()
                elif choice == '3':
                    self.create_suit()
                elif choice == '4':
                    self.show_items()
                elif choice == '5':
                    self.remove_item()
                elif choice == '6':
                    self.calculate_totals()
                elif choice == '7' and DOCX_AVAILABLE:
                    self.save_report()
                elif choice == '8':
                    self.show_history()
                elif choice == '9':
                    print("\nДо свидания!")
                    break
                else:
                    print("\nНеверный выбор! Пожалуйста, выберите 1-9")
                    self.wait_for_enter()

            except KeyboardInterrupt:
                print("\n\nПолучен сигнал прерывания. Завершение работы...")
                break
            except Exception as e:
                print(f"\nПроизошла ошибка: {e}")
                print("Продолжаем работу...")
                self.wait_for_enter()


if __name__ == "__main__":
    # Добавляем обработку сигналов для корректного завершения
    try:
        app = ClothingApp()
        app.run()
    except KeyboardInterrupt:
        print("\nПрограмма завершена пользователем.")
    except Exception as e:
        print(f"\nКритическая ошибка: {e}")
        sys.exit(1)